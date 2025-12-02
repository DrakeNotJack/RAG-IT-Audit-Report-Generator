from pathlib import Path
from typing import Dict, Any
from datetime import datetime
import json

from langchain_core.runnables import RunnablePassthrough, RunnableLambda, RunnableParallel
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.chat_models import ChatOllama


# ========= 1. Configuration =========
BASE_DIR = Path(__file__).resolve().parent.parent
CHROMA_DIR = BASE_DIR / "chroma_db"

EMBEDDING_MODEL = "nomic-embed-text"
LLM_MODEL = "qwen2.5:7b"

# Global LLM parameter configuration (centralized management)
LLM_PARAMS = {
    "parse": {  # Parameters used for parsing queries
        "temperature": 0.1,
        "max_tokens": 200,
        "top_p": 0.9
    },
    "generate": {  # Parameters used for report generation
        "temperature": 0.2,
        "max_tokens": 4000,
        "top_p": 0.9,
        "repetition_penalty": 1.1
    }
}

# ========= 2. Create LCEL chain =========
def create_audit_chain():
    """
    Create a chain for generating an ITGC audit report.
    Supports natural language input, with LLM-assisted parsing for client name and ITGC module.
    
    Input: {"query": str} (natural language audit inquiry)
    Output: Structured audit report as string
    """
    
    # 1. LLM-assisted parser: Extract client_name and module from natural language query
    def parse_natural_query(input_dict: Dict[str, Any]) -> Dict[str, Any]:
        """
        Parse natural language audit query to extract core structured info:
        - client_name: Name of the audited client (e.g., "ABC Manufacturing")
        - module: ITGC module name (e.g., "access_management", "change_management")
        Uses LLM for robust parsing (handles varied expressions).
        """
        natural_query = input_dict["query"].strip()
        
        # Prompt for LLM to parse structured info (JSON output required)
        parse_prompt = ChatPromptTemplate.from_template("""
        You are a structured data extractor for ITGC audit queries.
        
        EXTRACTION RULES:
        1. Extract ONLY these two fields:
        - client_name: Full name of the client/company
        - module: ITGC module name (must be: access_management, change_management, backup_management, incident_management)

        2. OUTPUT FORMAT: 
        - MUST be a plain JSON object WITHOUT any markdown formatting
        - NO ```json``` code blocks
        - NO additional text or explanations
        - Just the raw JSON: {{"client_name": "...", "module": "..."}}
        
        User's Audit Query: {natural_query}
        """)
        
        # Invoke LLM to get parsed result
        llm = ChatOllama(model=LLM_MODEL, **LLM_PARAMS["parse"])
        parse_result = llm.invoke(parse_prompt.format(natural_query=natural_query))
    
        # Debug / troubleshooting log (disabled by default)
        # print("=" * 80)
        # print("parse_result type:", type(parse_result))
        # print("\nFull parse_result:")
        # print(parse_result)
        # print("\nparse_result.content (LLM generated text):")
        # print(parse_result.content.strip())
        # print("=" * 80)
            
        # Parse JSON output (add fallback for edge cases)
        try:
            parsed_data = json.loads(parse_result.content.strip())
            # Validate required fields
            client_name = parsed_data.get("client_name", "Unknown Client")
            module = parsed_data.get("module", "Unknown module")
            # Standardize module format (replace spaces with underscores, lowercase)
            module = module.replace(" ", "_").lower()
        except (json.JSONDecodeError, AttributeError):
            # Fallback if LLM output is invalid
            client_name = "Unknown Client"
            module = "Unknown module"
        
        return {
            "client_name": client_name,
            "module": module,
            "natural_query": natural_query  # Preserve original query for context
        }
    
    # 2. Evidence retrieval: Fetch relevant docs from Chroma vector DB
    def retrieve_evidence(input_dict: Dict[str, Any]) -> Dict[str, Any]:
        """
        Retrieve relevant audit evidence from vector database based on parsed info:
        - ITGC standards specific to the module
        - Client's internal policies for the module
        - Past audit inquiry responses for the client/module
        """
        client_name = input_dict["client_name"]
        module = input_dict["module"]
        natural_query = input_dict["natural_query"]
        
        # Construct retrieval question (combine context + natural query for relevance)
        retrieval_question = f"""
        Audit Context: Client "{client_name}", ITGC Module: {module}
        User's Specific Audit Inquiry: {natural_query}
        
        Retrieval Requirements:
        1. Retrieve ONLY content directly relevant to {module} and the inquiry above
        2. Mandatory relevant materials:
           - ITGC standards specific to {module}
           - {client_name}'s official {module} policy documents
           - This client's past audit inquiry responses for {module}
        3. Content must support:
           - Control design effectiveness evaluation
           - Control gap identification
           - IT risk assessment
           - Audit conclusion formulation
        """.strip()
        
        # Initialize embeddings and vector DB
        embeddings = OllamaEmbeddings(model=EMBEDDING_MODEL,**LLM_PARAMS["generate"])
        vectordb = Chroma(
            embedding_function=embeddings,
            persist_directory=str(CHROMA_DIR),
        )
        
        # Perform similarity search (top 8 relevant docs)
        docs = vectordb.similarity_search(retrieval_question, k=8)
        
        # Build context string with metadata (for traceability)
        context = "\n\n---\n\n".join(
            f"[Metadata: {d.metadata}]\nContent: {d.page_content}"
            for d in docs
        )
        
        return {
            "client_name": client_name,
            "module": module,
            "context": context,
            "natural_query": natural_query
        }
    
    # 3. Prompt template for structured audit report
    prompt_template = ChatPromptTemplate.from_template("""
    You are an IT audit senior performing an ITGC review for
    Client: {client_name}
    ITGC Module: {module}
    User's Audit Inquiry: {natural_query}
    
    REQUIREMENTS:
    - Base ALL conclusions strictly on the provided retrieved evidence.
    - If insufficient evidence is found, explicitly state: "Insufficient evidence found for comprehensive audit assessment."
    - DO NOT invent or assume processes/tools not present in the evidence.
    - Use formal, precise audit language (English).
    - Follow the exact section structure below.
    
    RETRIEVED EVIDENCE (standards, policies, client practices):
    {context}
    
    REPORT STRUCTURE:
    1. Scope and Objective
       - Clearly state the audit scope (module + inquiry focus) and objectives.
    2. Process Description
       - Describe the client's actual processes/practices for the module (based solely on evidence).
    3. Design Assessment
       - Evaluate if control design complies with ITGC standards and client policies.
       - Cite evidence to support compliance/non-compliance conclusions.
    4. Potential Risks
       - Identify control gaps or weaknesses from the evidence.
       - Analyze associated risks (e.g., data breach, operational disruption).
    5. Recommendations
       - Provide specific, actionable remediation steps to address gaps/risks.
    6. Referenced Evidence
       - Briefly list the key supporting documents/chunks you relied on.
    """) #Requirement can be further polished for more accuracy, just a simple demo here, but 
    
    # 4. LLM and output parser initialization
    llm = ChatOllama(model=LLM_MODEL)
    output_parser = StrOutputParser()
    
    # 5. Full LCEL chain: Parse → Retrieve → Generate
    chain = (
        RunnablePassthrough()
        # Step 1: Parse the natural query to produce parsed_info (client_name, module, natural_query)
        | RunnableLambda(parse_natural_query)
        # Step 2: Run two branches in parallel: (1) retrieve evidence and generate report, (2) pass parsed_info through
        | RunnableParallel(
            # Branch 1: perform retrieval -> generate report
            report_chain=RunnableLambda(retrieve_evidence)
                        | prompt_template
                        | llm
                        | output_parser,
            # Branch 2: passthrough of parsed_info
            parsed_info=RunnablePassthrough()
        )
        # Step 3: Merge branch outputs into a single result dictionary
        | RunnableLambda(lambda x: {
            "report": x["report_chain"],  # final report text
            "client_name": x["parsed_info"]["client_name"],  # extracted from parsed_info
            "module": x["parsed_info"]["module"]  # extracted from parsed_info
        })
    )
    
    return chain


# ========= 3. Generate audit report =========
def generate_audit_report(natural_query: str) -> Dict[str, Any]:
    """
    Generate structured audit report from natural language query.
    
    Parameters:
        natural_query (str): Audit inquiry in natural language (e.g., "Audit ABC Manufacturing's access management...")
    
    Returns:
        Dict[str, Any]: Contains "report" (full report text), "client_name", "module"
    """
    # Initialize chain
    chain = create_audit_chain()
    
    # Invoke chain with natural language query
    parsed_input = {"query": natural_query}
    result = chain.invoke(parsed_input)
    
    return result

# ========= 4. Save report to Word document =========
def save_report_to_word(report_data: Dict[str, Any]) -> Path:
    """
    Save generated audit report to a Word document (docx format).
    
    Parameters:
        report_data (Dict[str, Any]): Output from generate_audit_report()
    
    Returns:
        Path: Full path to the saved Word document
    """
    from docx import Document
    
    # Extract data from report_data
    report_text = report_data["report"]
    client_name = report_data["client_name"]
    module = report_data["module"]
    
    # Create output directory (if not exists)
    reports_dir = BASE_DIR / "reports"
    reports_dir.mkdir(exist_ok=True)
    
    # Generate filename with timestamp (for uniqueness)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"ITGC_Audit_Report_{client_name}_{module}_{timestamp}.docx"
    file_path = reports_dir / filename
    
    # Create Word document
    doc = Document()
    
    # Add title and basic info
    doc.add_heading(f"ITGC Audit Report - {module.replace('_', ' ').title()}", 0)
    doc.add_paragraph(f"Client Name: {client_name}")
    doc.add_paragraph(f"Audit Module: {module.replace('_', ' ').title()}")
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()  # Add blank line
    
    # Add report content (split by paragraphs)
    doc.add_heading("Audit Findings & Conclusions", level=1)
    paragraphs = report_text.split("\n\n")
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    # Save document
    doc.save(file_path)
    print(f"Report saved successfully to: {file_path}")
    
    return file_path


# ========= 5. Main execution =========
if __name__ == "__main__":
    # Example: Natural language audit query (supports varied expressions)
    natural_query = """
    Could you audit the access management process of ABC Manufacturing? 
    Specifically, how do they revoke user access when employees leave the company?
    Is this process compliant with ISO 27001 Annex A.9.2.4? 
    Please identify any risks and provide improvement suggestions.
    """
    
    # Print header
    print("=" * 60)
    print("ITGC Audit Report Generator (Natural Language Support)")
    print("=" * 60)
    print(f"User's Audit Inquiry:\n{natural_query}\n")
    
    # Generate report
    print("Generating audit report... (this may take 10-60 seconds based on hardware)")
    report_data = generate_audit_report(natural_query)
    
    # Print generated report
    print("\n" + "=" * 60)
    print("Generated Audit Report")
    print("=" * 60)
    print(report_data["report"])
    
    
    # Can add if condition here, only if the report retrieve sufficient information then generate the report, otherwise output error message
    # Save to Word
    print("\n" + "=" * 60)
    save_report_to_word(report_data)
    
    print("\nProcess completed successfully!")